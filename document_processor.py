"""
ドキュメント処理モジュール
PDF、Word、Excelファイルからテキストを抽出します
"""
import os
import io
import zipfile
from typing import List, Dict
from pathlib import Path
import PyPDF2
from docx import Document
import openpyxl


class DocumentProcessor:
    """各種ドキュメント形式からテキストを抽出するクラス"""

    def __init__(self, documents_dir: str = "documents"):
        """
        Args:
            documents_dir: ドキュメントが格納されているディレクトリ
        """
        self.documents_dir = Path(documents_dir)

    def extract_text_from_pdf(self, file_path: Path) -> str:
        """PDFファイルからテキストを抽出"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"PDF読み込みエラー ({file_path}): {e}")
        return text

    def extract_text_from_word(self, file_path: Path) -> str:
        """Wordファイルからテキストを抽出（表の構造を保持、埋め込みExcel対応）"""
        text = ""
        try:
            doc = Document(file_path)

            # 段落と表を文書内の順序通りに処理
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    # 段落の処理
                    for paragraph in doc.paragraphs:
                        if paragraph._element == element:
                            if paragraph.text.strip():
                                text += paragraph.text + "\n"
                            break
                elif element.tag.endswith('tbl'):
                    # 表の処理
                    for table in doc.tables:
                        if table._tbl == element:
                            text += self._extract_table(table)
                            break

            # 埋め込みExcelファイルを抽出
            embedded_text = self._extract_embedded_excel(file_path)
            if embedded_text:
                text += "\n" + embedded_text

        except Exception as e:
            print(f"Word読み込みエラー ({file_path}): {e}")
        return text

    def _extract_embedded_excel(self, file_path: Path) -> str:
        """Word内に埋め込まれたExcelファイルを抽出"""
        text = ""
        try:
            with zipfile.ZipFile(file_path) as z:
                for name in z.namelist():
                    if name.startswith('word/embeddings/') and name.endswith('.xlsx'):
                        print(f"  埋め込みExcel発見: {name}")
                        with z.open(name) as excel_file:
                            excel_data = io.BytesIO(excel_file.read())
                            text += self._extract_excel_from_bytes(excel_data, name)
        except Exception as e:
            print(f"埋め込みExcel抽出エラー: {e}")
        return text

    def _extract_excel_from_bytes(self, excel_data: io.BytesIO, source_name: str) -> str:
        """BytesIOからExcelデータを抽出（部門ごとに分割）"""
        text = ""
        try:
            workbook = openpyxl.load_workbook(excel_data, data_only=True)
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]

                current_dept = None
                dept_data = []
                header_row = None

                for row in sheet.iter_rows(values_only=True):
                    cells = [str(cell) if cell is not None else "" for cell in row]
                    row_text = " | ".join(cells).strip()

                    if not any(c.strip() for c in cells):
                        continue

                    # 部門名の検出（「科」「部」「部門」「室」で終わる行）
                    first_cell = cells[0].strip().replace(" ", "")
                    is_dept_header = (
                        first_cell and
                        (first_cell.endswith('科') or first_cell.endswith('部') or
                         first_cell.endswith('部門') or first_cell.endswith('室') or
                         '看護' in first_cell or 'ステーション' in first_cell) and
                        len(first_cell) >= 2
                    )

                    # ヘッダー行の検出
                    if '始業時間' in row_text or '始業' in row_text and '終業' in row_text:
                        header_row = row_text
                        continue

                    if is_dept_header:
                        # 前の部門データを保存
                        if current_dept and dept_data:
                            text += f"\n【{current_dept}の勤務時間】\n"
                            if header_row:
                                text += f"{header_row}\n"
                            for d in dept_data:
                                text += f"{d}\n"
                            text += "\n"
                        # 新しい部門を開始
                        current_dept = first_cell
                        dept_data = []
                    elif current_dept and ('勤' in row_text or '番' in row_text or '曜' in row_text or ':' in row_text):
                        # 勤務データ行
                        dept_data.append(row_text)

                # 最後の部門データを保存
                if current_dept and dept_data:
                    text += f"\n【{current_dept}の勤務時間】\n"
                    if header_row:
                        text += f"{header_row}\n"
                    for d in dept_data:
                        text += f"{d}\n"
                    text += "\n"

        except Exception as e:
            print(f"埋め込みExcel読み込みエラー ({source_name}): {e}")
        return text

    def _extract_table(self, table) -> str:
        """表をマークダウン形式で抽出"""
        table_text = "\n【表】\n"

        for row_idx, row in enumerate(table.rows):
            # 各セルのテキストを取得（改行を削除してクリーン化）
            cells = []
            for cell in row.cells:
                cell_text = cell.text.replace('\n', ' ').strip()
                cells.append(cell_text)

            # 重複セル（結合セル）を除去
            unique_cells = []
            prev_text = None
            for cell_text in cells:
                if cell_text != prev_text:
                    unique_cells.append(cell_text)
                    prev_text = cell_text

            row_text = " | ".join(unique_cells)
            table_text += row_text + "\n"

            # ヘッダー行の後に区切り線を追加
            if row_idx == 0:
                table_text += "-" * 40 + "\n"

        table_text += "\n"
        return table_text

    def extract_text_from_excel(self, file_path: Path) -> str:
        """Excelファイルからテキストを抽出"""
        text = ""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"\n=== {sheet_name} ===\n"

                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text += row_text + "\n"
        except Exception as e:
            print(f"Excel読み込みエラー ({file_path}): {e}")
        return text

    def process_document(self, file_path: Path) -> Dict[str, str]:
        """
        ファイル形式に応じてテキストを抽出

        Returns:
            {'filename': str, 'content': str, 'file_type': str}
        """
        suffix = file_path.suffix.lower()

        if suffix == '.pdf':
            content = self.extract_text_from_pdf(file_path)
            file_type = 'PDF'
        elif suffix in ['.docx', '.doc']:
            content = self.extract_text_from_word(file_path)
            file_type = 'Word'
        elif suffix in ['.xlsx', '.xls']:
            content = self.extract_text_from_excel(file_path)
            file_type = 'Excel'
        else:
            content = ""
            file_type = 'Unknown'

        return {
            'filename': file_path.name,
            'content': content,
            'file_type': file_type,
            'file_path': str(file_path)
        }

    def process_all_documents(self) -> List[Dict[str, str]]:
        """
        ドキュメントディレクトリ内のすべてのファイルを処理

        Returns:
            処理されたドキュメントのリスト
        """
        documents = []
        supported_extensions = ['.pdf', '.docx', '.doc', '.xlsx', '.xls']

        if not self.documents_dir.exists():
            print(f"警告: ディレクトリが存在しません: {self.documents_dir}")
            return documents

        for file_path in self.documents_dir.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                print(f"処理中: {file_path.name}")
                doc = self.process_document(file_path)
                if doc['content'].strip():
                    documents.append(doc)
                else:
                    print(f"警告: テキストが抽出できませんでした: {file_path.name}")

        return documents
