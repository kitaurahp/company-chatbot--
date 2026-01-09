"""
テスト用のサンプル社内規定ファイルを作成
"""
from docx import Document
import openpyxl

# Word文書を作成（就業規則サンプル）
def create_employment_rules():
    doc = Document()
    doc.add_heading('就業規則', 0)

    doc.add_heading('第1章 総則', level=1)
    doc.add_paragraph('第1条（目的）')
    doc.add_paragraph('この規則は、株式会社〇〇における従業員の就業に関する事項を定めることを目的とする。')

    doc.add_paragraph('第2条（適用範囲）')
    doc.add_paragraph('この規則は、当社の正社員、契約社員、パート社員に適用する。')

    doc.add_heading('第2章 勤務時間', level=1)
    doc.add_paragraph('第3条（勤務時間）')
    doc.add_paragraph('1. 所定労働時間は、1日8時間、週40時間とする。')
    doc.add_paragraph('2. 始業時刻は午前9時、終業時刻は午後6時とする。')
    doc.add_paragraph('3. 休憩時間は正午から午後1時までの1時間とする。')

    doc.add_paragraph('第4条（休日）')
    doc.add_paragraph('1. 休日は、土曜日、日曜日及び国民の祝日とする。')
    doc.add_paragraph('2. 年末年始（12月29日から1月3日）は休日とする。')

    doc.add_heading('第3章 休暇', level=1)
    doc.add_paragraph('第5条（年次有給休暇）')
    doc.add_paragraph('1. 入社6ヶ月経過後、10日の年次有給休暇を付与する。')
    doc.add_paragraph('2. 以降、1年ごとに付与日数を増加する。')
    doc.add_paragraph('3. 有給休暇の申請は、原則として3日前までに所属長に届け出ること。')

    doc.add_paragraph('第6条（特別休暇）')
    doc.add_paragraph('1. 慶弔休暇：結婚の場合5日、配偶者の出産2日、忌引き（続柄により3〜7日）')
    doc.add_paragraph('2. 育児休暇：子が1歳に達するまで取得可能')
    doc.add_paragraph('3. 介護休暇：要介護状態の家族1人につき年5日まで')

    doc.add_heading('第4章 リモートワーク', level=1)
    doc.add_paragraph('第7条（リモートワーク）')
    doc.add_paragraph('1. 従業員は週2日までリモートワークを申請できる。')
    doc.add_paragraph('2. リモートワークの申請は前日までに所属長の承認を得ること。')
    doc.add_paragraph('3. リモートワーク中も所定の勤務時間を遵守すること。')

    doc.save('documents/就業規則.docx')
    print('✓ 就業規則.docx を作成しました')

# Excel文書を作成（旅費規定サンプル）
def create_travel_expenses():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '旅費規定'

    # タイトル
    ws['A1'] = '旅費規定'
    ws['A1'].font = openpyxl.styles.Font(size=16, bold=True)

    # 交通費
    ws['A3'] = '1. 交通費'
    ws['A4'] = '種別'
    ws['B4'] = '支給基準'
    ws['A5'] = '電車・バス'
    ws['B5'] = '実費（領収書必須）'
    ws['A6'] = '新幹線'
    ws['B6'] = '指定席可（片道2時間以上の場合）'
    ws['A7'] = '航空機'
    ws['B7'] = 'エコノミークラス（片道3時間以上の場合）'
    ws['A8'] = 'タクシー'
    ws['B8'] = '原則不可（終電後、荷物多数の場合は可）'

    # 宿泊費
    ws['A10'] = '2. 宿泊費'
    ws['A11'] = '地域'
    ws['B11'] = '上限額（1泊）'
    ws['A12'] = '東京23区内'
    ws['B12'] = '12,000円'
    ws['A13'] = '大阪市内'
    ws['B13'] = '10,000円'
    ws['A14'] = 'その他地域'
    ws['B14'] = '8,000円'
    ws['A15'] = '※朝食込みの金額'

    # 日当
    ws['A17'] = '3. 日当'
    ws['A18'] = '出張日数'
    ws['B18'] = '日当'
    ws['A19'] = '日帰り'
    ws['B19'] = '1,000円'
    ws['A20'] = '1泊2日以上'
    ws['B20'] = '2,000円/日'

    # 申請方法
    ws['A22'] = '4. 申請方法'
    ws['A23'] = '・出張前：出張申請書を所属長に提出し承認を得る'
    ws['A24'] = '・出張後：旅費精算書に領収書を添付して経理部に提出'
    ws['A25'] = '・精算期限：出張終了後7日以内'

    # 列幅調整
    ws.column_dimensions['A'].width = 20
    ws.column_dimensions['B'].width = 40

    wb.save('documents/旅費規定.xlsx')
    print('✓ 旅費規定.xlsx を作成しました')

# 経費精算ルールを作成
def create_expense_rules():
    doc = Document()
    doc.add_heading('経費精算ルール', 0)

    doc.add_heading('1. 基本原則', level=1)
    doc.add_paragraph('業務上必要な経費は、事前申請または事後精算により会社が負担します。')
    doc.add_paragraph('私的な支出は経費として認められません。')

    doc.add_heading('2. 経費として認められる項目', level=1)
    doc.add_paragraph('・交通費（通勤定期、出張交通費）')
    doc.add_paragraph('・宿泊費（出張時）')
    doc.add_paragraph('・会議費（社内外の会議での飲食費）')
    doc.add_paragraph('・接待交際費（取引先との接待、手土産等）')
    doc.add_paragraph('・消耗品費（文房具、PC周辺機器等）')
    doc.add_paragraph('・書籍費（業務関連書籍、技術書等）')
    doc.add_paragraph('・研修費（セミナー参加費、資格取得費用等）')

    doc.add_heading('3. 申請・精算手順', level=1)
    doc.add_paragraph('【事前申請が必要な経費】')
    doc.add_paragraph('・10,000円以上の支出')
    doc.add_paragraph('・接待交際費')
    doc.add_paragraph('・研修参加費')

    doc.add_paragraph('【精算方法】')
    doc.add_paragraph('1. 経費精算システムに必要事項を入力')
    doc.add_paragraph('2. 領収書の写真をアップロード')
    doc.add_paragraph('3. 所属長の承認を得る')
    doc.add_paragraph('4. 経理部で確認後、翌月15日に振込')

    doc.add_heading('4. 領収書について', level=1)
    doc.add_paragraph('・必ず宛名が会社名（株式会社〇〇）の領収書を取得すること')
    doc.add_paragraph('・レシートでも可（ただし、品目が明確なもの）')
    doc.add_paragraph('・領収書は原本を保管し、精算後も3年間保管すること')

    doc.add_heading('5. 注意事項', level=1)
    doc.add_paragraph('・虚偽の申請は懲戒処分の対象となります')
    doc.add_paragraph('・個人のクレジットカードで支払った場合も精算可能です')
    doc.add_paragraph('・不明点は経理部に問い合わせてください')

    doc.save('documents/経費精算ルール.docx')
    print('✓ 経費精算ルール.docx を作成しました')

if __name__ == '__main__':
    print('サンプル文書を作成中...\n')
    create_employment_rules()
    create_travel_expenses()
    create_expense_rules()
    print('\n✓ すべてのサンプル文書を作成しました！')
    print('documentsフォルダを確認してください。')
